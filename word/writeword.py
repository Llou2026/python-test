from docx import Document
from docx.shared import Pt,RGBColor
from docx.enum.style import WD_STYLE_TYPE
#1.创建一个文档对象
document= Document()#新建文档对象
#Document('info.docx')#读取现有的word建立文档对象
#2.写入内容
document.add_heading('慕课网简介',level=4)#写入标题,第二个参数level标题的级别取值是0-9
#样式
style=document.styles.add_style('textstyle',WD_STYLE_TYPE.PARAGRAPH)
#WD_STYLE_TYPE.CHARACTER → 字符样式（只给选中的一段文字生效）
#WD_STYLE_TYPE.TABLE → 表格样式
#WD_STYLE_TYPE.LIST → 列表样式
print(style.style_id)
print(style.name)
style.font.size=Pt(5)
#删除样式
document.styles['textstyle'].delete()
#写入段落
p1=document.add_paragraph('从OpenClaw的缔造者、如今在OpenAI主导下一代个人智能体的Peter Steinberger，到Claude Code的核心开发者Boris Cherny，这些站在浪潮之巅的工程师们，正在集体转向一种全新的工作范式：循环（loops）',style='textstyle')
p1.insert_paragraph_before('！！！！！！！！！！！在段落前插入一个新的段落')#通过p1在段落前插入一个新的段落
format=p1.paragraph_format
#左右的一个缩进
format.left_indent=Pt(20)#表示左侧的一个缩进
format.right_indent=Pt(20)#表示右侧的一个缩进
format.first_line_indent=Pt(20)#表示首行的一个缩进
#行间距
format.line_spacing=1.5#设置行间距
run=p1.add_run('！！！！！！！！！！！！！！！追加的内容，看看我是不是在末尾')
#单独设置上面这一行的字体、字号以及文字颜色
run.font.size=Pt(12)#设置字体大小，单位是磅
run.font.name='微软雅黑'#设置字体
run.font.color.rgb=RGBColor(235,33,24)
run1=p1.add_run('-----------再次用p1.add_run追加你看我会在那个位置出现------------')
#加粗、下划线、斜体
run1.bold=True #加粗
run1.font.underline=True#下划线
run1.font.italic=True#斜体
#插入图片
document.add_picture('logo.jpg')
document.add_picture('logo.jpg',Pt(20),Pt(30))#文件名，宽度，高度
#插入表格
table=document.add_table(rows=1,cols=3,style='Medium List 2')#新建一个表格，定义table接收
header_cells=table.rows[0].cells#通过rows得到所有行，通过[0]得到第一行
header_cells[0].text='月份'
header_cells[1].text='预期销售额'
header_cells[2].text='实际销售额'
#数据
data=[
    ['一月份',500,600],
    ['二月份',1000,400],
    ['三月份',1003,900]
]
for item in data:
    rows_cells=table.add_row().cells
    #add_row()在表格底部增加一行，cells表示获取到刚刚增加的一行的单元格对象
    rows_cells[0].text=item[0]
    rows_cells[1].text=str(item[1])
    rows_cells[2].text=str(item[2])
    #rows_cells接收的是一个元组，不能直接改变，通过text属性修改
    
    #获取表格,tables是获取到文档中所有的表格，通过索引可以访问到具体某个表格
    #print(len(document.tables[0].rows))#刚刚创建的表格，通过rows打印总行数
#3.保存文档
document.save('info.docx')#传入保存的文件名